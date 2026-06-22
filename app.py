import streamlit as st
import google.generativeai as genai
import pypdf
import json
import requests
import hashlib
import pandas as pd
from datetime import datetime
from docx import Document
from io import BytesIO

# ---------------- CONFIG ----------------
st.set_page_config(page_title="AI Dental Research Hub", layout="wide")

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
SHEETS_URL = st.secrets["SHEETS_WEBAPP_URL"]

MODEL = "gemini-1.5-flash"

# ---------------- USERS ----------------
USERS = {
    "resident1": "1234",
    "resident2": "1234",
    "admin": "admin"
}

if "user" not in st.session_state:
    st.session_state.user = None

# ---------------- LOGIN ----------------
if not st.session_state.user:

    st.title("🔐 Login")

    u = st.text_input("User")
    p = st.text_input("Password", type="password")

    if st.button("Login"):
        if u in USERS and USERS[u] == p:
            st.session_state.user = u
            st.rerun()
        else:
            st.error("Wrong credentials")

    st.stop()

user = st.session_state.user

# ---------------- LOAD DATA ----------------
try:
    db = requests.get(SHEETS_URL, timeout=10).json()
except:
    db = {"data": []}

# flatten
all_rows = []
for j in db.values():
    if isinstance(j, list):
        all_rows.extend(j[1:])

df = pd.DataFrame(all_rows, columns=[
    "user","journal","month","title","summary","one_liner","topic","tags"
])

# ---------------- UTIL ----------------
def hash_text(text):
    return hashlib.sha256(text.encode()).hexdigest()

# ---------------- TAGGING ----------------
def auto_tag(text):
    model = genai.GenerativeModel(MODEL)

    prompt = f"""
Classify this dental paper into tags.
Return JSON array only.

Tags allowed:
implantology, periodontology, regeneration, peri-implantitis, systemic, microbiology, surgery

TEXT:
{text[:8000]}
"""

    res = model.generate_content(prompt)
    return json.loads(res.text.replace("```",""))

# ---------------- SUMMARIZE ----------------
def summarize(text):

    model = genai.GenerativeModel(MODEL)

    prompt = f"""
Return JSON ONLY:

- summary (6 bullet points Hebrew)
- one_liner (Hebrew)
- topic (Hebrew)
- title_and_authors

TEXT:
{text[:15000]}
"""

    res = model.generate_content(prompt)
    return json.loads(res.text.replace("```",""))

# ---------------- EXPORT WORD ----------------
def to_word(rows):

    doc = Document()
    doc.add_heading("Dental Papers Export", 0)

    for r in rows:
        doc.add_heading(r["title"], level=1)
        doc.add_paragraph("Summary:")
        doc.add_paragraph(r["summary"])
        doc.add_paragraph("Takeaway:")
        doc.add_paragraph(r["one_liner"])

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    return bio

# ---------------- UI ----------------
st.title(f"🦷 Dental Research Hub - {user}")

tab1, tab2, tab3, tab4 = st.tabs([
    "📤 Upload",
    "🔍 Search",
    "📊 Archive",
    "⬇️ Export"
])

# ---------------- UPLOAD ----------------
with tab1:

    file = st.file_uploader("Upload PDF", type=["pdf"])

    if file and st.button("Process"):

        pdf = pypdf.PdfReader(file)

        text = ""
        for p in pdf.pages[:10]:
            t = p.extract_text()
            if t:
                text += t

        doc_id = hash_text(text)

        # tagging + summary
        summary = summarize(text)
        tags = auto_tag(text)

        payload = {
            "user": user,
            "journal": "AUTO",
            "month": datetime.now().strftime("%m.%Y"),
            "title": summary.get("title_and_authors",""),
            "summary": summary.get("summary",""),
            "one_liner": summary.get("one_liner",""),
            "topic": summary.get("topic",""),
            "tags": ",".join(tags) if isinstance(tags, list) else str(tags)
        }

        r = requests.post(SHEETS_URL, json=payload)

        if r.status_code == 200:
            st.success("Saved")
        else:
            st.error(r.text)

# ---------------- SEARCH ----------------
with tab2:

    q = st.text_input("Search papers (title / tag / topic)")

    if q:
        results = df[
            df.astype(str).apply(lambda x: x.str.contains(q, case=False).any(), axis=1)
        ]

        st.write(results)

# ---------------- ARCHIVE ----------------
with tab3:

    user_df = df[df["user"] == user]

    st.dataframe(user_df)

# ---------------- EXPORT ----------------
with tab4:

    user_df = df[df["user"] == user].to_dict("records")

    if st.button("Export to Word"):

        file = to_word(user_df)

        st.download_button(
            "Download Word",
            file,
            file_name="papers.docx"
        )

    if st.button("Export to Excel"):

        excel = BytesIO()
        pd.DataFrame(user_df).to_excel(excel, index=False)
        excel.seek(0)

        st.download_button(
            "Download Excel",
            excel,
            file_name="papers.xlsx"
        )
